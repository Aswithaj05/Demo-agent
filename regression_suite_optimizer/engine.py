"""Document parsing, requirement extraction, and regression test case generation.

Two generation modes:
- AI mode: uses Claude (requires ANTHROPIC_API_KEY / ANTHROPIC_AUTH_TOKEN).
- Demo mode: deterministic rule-based generator, works fully offline.
"""

from __future__ import annotations

import io
import json
import os
import re
import zipfile
import zlib
import xml.etree.ElementTree as ET

SS_NS = "{http://schemas.openxmlformats.org/spreadsheetml/2006/main}"

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".txt"}


# ---------------------------------------------------------------------------
# Document text extraction
# ---------------------------------------------------------------------------

def extract_text(filename: str, data: bytes) -> str:
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".txt":
        return data.decode("utf-8", errors="replace")
    if ext == ".docx":
        return _docx_text(data)
    if ext == ".xlsx":
        return _xlsx_text(data)
    if ext == ".pdf":
        return _pdf_text(data)
    raise ValueError(f"Unsupported file type: {ext}")


def _docx_text(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append("\t".join(cells))
    return "\n".join(parts)


def _xlsx_text(data: bytes) -> str:
    zf = zipfile.ZipFile(io.BytesIO(data))
    shared: list[str] = []
    if "xl/sharedStrings.xml" in zf.namelist():
        root = ET.fromstring(zf.read("xl/sharedStrings.xml"))
        for si in root.findall(f"{SS_NS}si"):
            shared.append("".join(t.text or "" for t in si.iter(f"{SS_NS}t")))
    lines: list[str] = []
    for name in sorted(zf.namelist()):
        if not re.match(r"xl/worksheets/sheet\d+\.xml$", name):
            continue
        root = ET.fromstring(zf.read(name))
        for row in root.iter(f"{SS_NS}row"):
            values: list[str] = []
            for cell in row.iter(f"{SS_NS}c"):
                ctype = cell.get("t")
                if ctype == "s":
                    v = cell.find(f"{SS_NS}v")
                    idx = int(v.text) if v is not None and v.text else -1
                    values.append(shared[idx] if 0 <= idx < len(shared) else "")
                elif ctype == "inlineStr":
                    values.append("".join(t.text or "" for t in cell.iter(f"{SS_NS}t")))
                else:
                    v = cell.find(f"{SS_NS}v")
                    values.append(v.text or "" if v is not None else "")
            if any(v.strip() for v in values):
                lines.append("\t".join(values))
    return "\n".join(lines)


def _pdf_text(data: bytes) -> str:
    """PDF text extraction: pypdf if installed, else the built-in extractor."""
    try:
        import pypdf  # optional — used automatically when installed

        reader = pypdf.PdfReader(io.BytesIO(data))
        text = "\n".join((page.extract_text() or "") for page in reader.pages)
        if text.strip():
            return text
    except ImportError:
        pass
    except Exception:  # noqa: BLE001 — fall through to the built-in extractor
        pass
    return _pdf_text_builtin(data)


def _pdf_decode_stream(chunk: bytes) -> bytes:
    """Decode a raw PDF stream body: Flate, ASCII85(+Flate), ASCIIHex, or raw."""
    import base64

    chunk = chunk.strip(b"\r\n")
    try:
        return zlib.decompress(chunk)
    except zlib.error:
        pass
    # ASCII85 (e.g. PDFsharp output), optionally wrapping Flate
    candidate = re.sub(rb"\s+", b"", chunk)
    if candidate.endswith(b"~>"):
        candidate = candidate[:-2]
    if candidate and re.fullmatch(rb"[!-uz]+", candidate):
        try:
            decoded = base64.a85decode(candidate)
            try:
                return zlib.decompress(decoded)
            except zlib.error:
                return decoded
        except ValueError:
            pass
    # ASCIIHex
    hex_candidate = re.sub(rb"\s+", b"", chunk).rstrip(b">")
    if hex_candidate and re.fullmatch(rb"[0-9A-Fa-f]+", hex_candidate):
        try:
            decoded = bytes.fromhex(hex_candidate.decode("ascii"))
            try:
                return zlib.decompress(decoded)
            except zlib.error:
                return decoded
        except ValueError:
            pass
    return chunk


_PDF_TOKEN = re.compile(
    r"\((?:\\.|[^\\()])*\)"      # literal string
    r"|<[0-9A-Fa-f\s]+>"          # hex string
    r"|\bT[dD*]\b|\bTm\b|\bET\b"  # line-positioning operators -> newline
)


def _pdf_literal(token: str) -> str:
    body = token[1:-1]
    body = re.sub(r"\\(\d{1,3})", lambda m: chr(int(m.group(1), 8) & 0xFF), body)
    return (
        body.replace("\\n", "\n").replace("\\r", "").replace("\\t", " ")
        .replace("\\(", "(").replace("\\)", ")").replace("\\\\", "\\")
    )


def _pdf_hex_string(token: str) -> str:
    digits = re.sub(r"\s+", "", token[1:-1])
    if len(digits) % 2:
        digits += "0"
    try:
        raw = bytes.fromhex(digits)
    except ValueError:
        return ""
    # Heuristic: UTF-16BE (2-byte CID/Identity-H) vs single-byte encoding
    if len(raw) >= 2 and raw[0] == 0 and len(raw) % 2 == 0:
        try:
            return raw.decode("utf-16-be")
        except UnicodeDecodeError:
            pass
    return raw.decode("latin-1", errors="replace")


def _pdf_text_builtin(data: bytes) -> str:
    """Best-effort stdlib PDF text extraction (Flate/ASCII85/ASCIIHex + Tj/TJ)."""
    parts: list[str] = []
    for match in re.finditer(rb"stream\r?\n(.*?)endstream", data, re.S):
        chunk = _pdf_decode_stream(match.group(1))
        try:
            text = chunk.decode("latin-1")
        except UnicodeDecodeError:
            continue
        if "Tj" not in text and "TJ" not in text:
            continue
        for token_match in _PDF_TOKEN.finditer(text):
            token = token_match.group(0)
            if token.startswith("("):
                parts.append(_pdf_literal(token))
            elif token.startswith("<"):
                parts.append(_pdf_hex_string(token))
            else:  # Td / TD / T* / Tm / ET — new line of text
                if parts and parts[-1] != "\n":
                    parts.append("\n")
        parts.append("\n")
    result = "".join(parts).strip()
    if not result:
        raise ValueError(
            "Could not extract text from this PDF (it may be scanned/image-only "
            "or use an unsupported encoding). Install pypdf for broader support "
            "(python -m pip install pypdf), or convert it to DOCX/TXT and retry."
        )
    return result


# ---------------------------------------------------------------------------
# Requirement extraction
# ---------------------------------------------------------------------------

REQ_VERBS = re.compile(
    r"\b(shall|must|should|will|needs? to|able to|allow|enable|support|validate|"
    r"verify|display|generate|process|calculate|send|store|reject|prevent|"
    r"restrict|require[sd]?|provide)\b",
    re.I,
)
REQ_ID = re.compile(r"^\s*(?:REQ|FR|BR|US|NFR)[-_ ]?\d+", re.I)

MODULE_KEYWORDS = [
    (re.compile(r"\b(login|logout|password|authenticat|sign[- ]?in|otp|mfa|2fa|session)\b", re.I), "Authentication"),
    (re.compile(r"\b(payment|invoice|billing|refund|transaction|checkout|card)\b", re.I), "Payments"),
    (re.compile(r"\b(report|dashboard|analytic|chart|export|summary)\b", re.I), "Reporting"),
    (re.compile(r"\b(user|profile|account|role|permission|admin)\b", re.I), "User Management"),
    (re.compile(r"\b(search|filter|sort|query)\b", re.I), "Search"),
    (re.compile(r"\b(upload|download|document|file|attachment)\b", re.I), "Documents"),
    (re.compile(r"\b(api|endpoint|integration|webhook|service)\b", re.I), "API"),
    (re.compile(r"\b(email|notification|alert|sms|message)\b", re.I), "Notifications"),
    (re.compile(r"\b(order|cart|product|inventory|catalog|shipment)\b", re.I), "Orders"),
]

HIGH_PRIORITY = re.compile(
    r"\b(payment|security|password|authenticat|compliance|audit|data loss|"
    r"critical|mandatory|shall|must|refund|transaction|privacy|encrypt)\b", re.I
)
LOW_PRIORITY = re.compile(r"\b(color|colour|tooltip|label|font|icon|cosmetic|alignment)\b", re.I)

NFR_PATTERN = re.compile(
    r"\b(performance|response time|latency|throughput|scalab|availab|reliab|uptime|"
    r"security|encrypt|complian|audit trail|usab|accessib|maintainab|portab|recoverab|"
    r"concurrent users?|load of|sla)\b", re.I
)
BUSINESS_RULE_PATTERN = re.compile(
    r"\b(rule|policy|calculat|eligib|approv|discount|tax|fee|interest|threshold|"
    r"rate|limit of)\b", re.I
)
_ALLOW_WORDS = re.compile(r"\b(allow|enable|permit|grant)\b", re.I)
_DENY_WORDS = re.compile(r"\b(restrict|prevent|reject|disallow|deny|block)\b", re.I)


def _is_heading(line: str) -> bool:
    stripped = line.strip()
    if not stripped or len(stripped) > 60 or stripped.endswith("."):
        return False
    if stripped.endswith(":"):
        return True
    if re.match(r"^\d+(\.\d+)*\.?\s+\S", stripped) and not REQ_VERBS.search(stripped):
        return True
    if stripped.isupper() and len(stripped) > 3:
        return True
    words = stripped.split()
    if 1 < len(words) <= 6 and all(w[:1].isupper() or not w[:1].isalpha() for w in words):
        return not REQ_VERBS.search(stripped)
    return False


def _infer_module(text: str, current_heading: str) -> str:
    for pattern, module in MODULE_KEYWORDS:
        if pattern.search(text):
            return module
    return current_heading or "General"


def extract_requirements(doc_texts: dict[str, str]) -> list[dict]:
    """Extract requirement statements from the combined document text."""
    requirements: list[dict] = []
    seen: set[str] = set()
    counter = 0
    for source, text in doc_texts.items():
        heading = ""
        for raw_line in text.splitlines():
            line = re.sub(r"\s+", " ", raw_line).strip(" \t-*•")
            if not line:
                continue
            if _is_heading(line):
                heading = line.rstrip(":").strip()
                heading = re.sub(r"^\d+(\.\d+)*\.?\s*", "", heading)
                continue
            in_req_section = bool(re.search(r"requirement|user stor|acceptance", heading, re.I))
            is_req = (
                bool(REQ_ID.match(line))
                or (len(line) >= 25 and REQ_VERBS.search(line))
                or (in_req_section and 12 <= len(line) <= 300)
            )
            if not is_req:
                continue
            norm = re.sub(r"[^a-z0-9]", "", line.lower())
            if norm in seen:
                continue
            seen.add(norm)
            counter += 1
            requirements.append({
                "id": f"REQ-{counter:03d}",
                "text": line[:500],
                "module": _infer_module(line, heading),
                "source": source,
                "type": "Non-Functional" if NFR_PATTERN.search(line) else "Functional",
            })
    return requirements


def extract_business_rules(requirements: list[dict]) -> list[dict]:
    """Business-rule-bearing requirements (pricing, eligibility, approval, limits, ...)."""
    rules: list[dict] = []
    for req in requirements:
        if BUSINESS_RULE_PATTERN.search(req["text"]):
            rules.append({
                "id": f"BR-{len(rules) + 1:03d}",
                "text": req["text"],
                "module": req["module"],
                "requirement_ref": req["id"],
                "source": req["source"],
            })
    return rules


def find_business_rule_conflicts(rules: list[dict]) -> list[dict]:
    """Heuristic: two business rules in the same module using opposing allow/deny language."""
    conflicts: list[dict] = []
    by_module: dict[str, list[dict]] = {}
    for r in rules:
        by_module.setdefault(r["module"], []).append(r)
    for module, rs in by_module.items():
        for i in range(len(rs)):
            for j in range(i + 1, len(rs)):
                a, b = rs[i], rs[j]
                if (_ALLOW_WORDS.search(a["text"]) and _DENY_WORDS.search(b["text"])) or \
                   (_ALLOW_WORDS.search(b["text"]) and _DENY_WORDS.search(a["text"])):
                    conflicts.append({
                        "a": a["id"], "b": b["id"], "module": module,
                        "note": f"{a['id']} and {b['id']} both govern {module} but use opposing "
                                f"allow/deny language — verify they don't contradict.",
                    })
    return conflicts


# ---------------------------------------------------------------------------
# Demo-mode (rule-based) test case generation
# ---------------------------------------------------------------------------

def _priority_for(text: str) -> str:
    if HIGH_PRIORITY.search(text):
        return "High"
    if LOW_PRIORITY.search(text):
        return "Low"
    return "Medium"


def _short(text: str, n: int = 110) -> str:
    return text if len(text) <= n else text[: n - 3] + "..."


CRITICAL_SEVERITY_PATTERN = re.compile(
    r"\b(payment|security|password|authenticat|compliance|audit|data loss|"
    r"privacy|encrypt|refund|transaction|pci|gdpr)\b", re.I
)
_HIGH_IMPACT_CATEGORIES = {"Integration", "End-to-End", "Business Rule Validation", "API Validation"}


def _severity_for(text: str, priority: str) -> str:
    if CRITICAL_SEVERITY_PATTERN.search(text):
        return "Critical"
    return priority


def _regression_impact_for(category: str, priority: str) -> str:
    if category in _HIGH_IMPACT_CATEGORIES or priority == "High":
        return "High"
    return "Medium" if priority == "Medium" else "Low"


def _automation_candidate_for(category: str) -> str:
    return "No" if category == "UI Validation" else "Yes"


def _test_objective_for(category: str, module: str) -> str:
    return f"Validate {category.lower()} behavior for the '{module}' module against the stated requirement."


class _TCBuilder:
    def __init__(self) -> None:
        self.cases: list[dict] = []
        self.seen_scenarios: set[str] = set()
        self.duplicates_skipped = 0
        self.duplicate_examples: list[str] = []
        self._n = 0

    def add(self, module: str, req: dict | None, scenario: str, category: str,
            steps: list[str], expected: str, priority: str,
            preconditions: str = "Application is accessible and test user is available.",
            test_data: str = "As specified in the linked requirement.") -> None:
        norm = re.sub(r"[^a-z0-9]", "", scenario.lower())
        if norm in self.seen_scenarios:
            self.duplicates_skipped += 1
            if len(self.duplicate_examples) < 25:
                self.duplicate_examples.append(scenario)
            return
        self.seen_scenarios.add(norm)
        self._n += 1
        self.cases.append({
            "id": f"TC-{self._n:03d}",
            "module": module,
            "requirement_ref": req["id"] if req else "N/A",
            "requirement_description": req["text"] if req else "Cross-module / suite-level check — not tied to a single requirement.",
            "user_story_ref": req["id"].replace("REQ", "US") if req else "N/A",
            "scenario": scenario,
            "test_objective": _test_objective_for(category, module),
            "preconditions": preconditions,
            "steps": steps,
            "test_data": test_data,
            "expected_result": expected,
            "priority": priority,
            "severity": _severity_for(req["text"] if req else scenario, priority),
            "category": category,
            "regression_impact": _regression_impact_for(category, priority),
            "automation_candidate": _automation_candidate_for(category),
            "traceability": f"{req['id']} ({req['source']})" if req else "Cross-module",
        })


VAGUE_LANGUAGE_PATTERN = re.compile(
    r"\b(tbd|appropriate|reasonable|etc\.?|and so on|approximately|as needed|if applicable|"
    r"where necessary|sufficient|acceptable)\b", re.I
)


def _extract_assumptions(requirements: list[dict]) -> list[str]:
    """Flag requirements using vague/underspecified language (demo engine only —
    no real language understanding, so this is a lightweight heuristic, not analysis)."""
    assumptions: list[str] = []
    for req in requirements:
        match = VAGUE_LANGUAGE_PATTERN.search(req["text"])
        if match:
            assumptions.append(
                f"{req['id']} uses vague language ('{match.group(0)}') — assumed standard "
                f"enterprise behavior applies: \"{_short(req['text'], 140)}\""
            )
    return assumptions


def demo_generate(requirements: list[dict]) -> dict:
    """Deterministic rule-based regression suite generation (no API needed)."""
    b = _TCBuilder()

    for req in requirements:
        text, module = req["text"], req["module"]
        prio = _priority_for(text)
        subject = _short(text)

        b.add(module, req,
              f"Verify that: {subject}",
              "Functional Regression",
              ["Log in with a valid user for the module.",
               f"Navigate to the {module} area.",
               "Perform the action described in the requirement with valid inputs.",
               "Observe the system behavior and any confirmation shown."],
              "The system behaves exactly as stated in the requirement; correct data is persisted/displayed.",
              prio)

        b.add(module, req,
              f"Negative — attempt the operation with invalid input/state for: {_short(text, 80)}",
              "Negative Scenario",
              [f"Navigate to the {module} area.",
               "Attempt the operation with invalid, missing, or unauthorized input.",
               "Observe the error handling."],
              "The system rejects the operation with a clear, user-friendly error message; no partial data is saved.",
              prio, test_data="Invalid/malformed values; unauthorized user where applicable.")

        if re.search(r"\b(field|form|input|mandatory|required|format|valid)\b", text, re.I):
            b.add(module, req,
                  f"Validation — mandatory fields and formats for: {_short(text, 80)}",
                  "Data Validation",
                  ["Open the relevant form/screen.",
                   "Submit with each mandatory field empty, one at a time.",
                   "Enter values violating expected formats (email, date, numeric).",
                   "Correct all values and submit."],
                  "Each violation produces a specific inline validation message; valid submission succeeds.",
                  prio, test_data="Empty strings, wrong formats, oversized values.")

        if re.search(r"\d|\b(limit|max|min|range|length|between|up to|at least)\b", text, re.I):
            b.add(module, req,
                  f"Boundary values for: {_short(text, 85)}",
                  "Boundary Value",
                  ["Identify the numeric/length limits in the requirement.",
                   "Test at limit-1, limit, and limit+1.",
                   "Verify behavior at each boundary."],
                  "Values at and below the limit are accepted; values beyond the limit are rejected with a message.",
                  prio, test_data="limit-1 / limit / limit+1 for each stated bound.")

        if re.search(r"\b(error|exception|fail|invalid|reject|timeout|unavailable)\b", text, re.I):
            b.add(module, req,
                  f"Error handling — system failure during: {_short(text, 80)}",
                  "Error Handling / Exception",
                  ["Begin the operation normally.",
                   "Simulate a failure (dependency down, timeout, network drop).",
                   "Retry after restoring the dependency."],
                  "A graceful error is shown, no data corruption occurs, and the retry succeeds.",
                  "High")

        if re.search(r"\b(rule|policy|calculat|eligib|approv|discount|tax|fee|interest)\b", text, re.I):
            b.add(module, req,
                  f"Business rule check for: {_short(text, 85)}",
                  "Business Rule Validation",
                  ["Prepare data sets that satisfy and violate the business rule.",
                   "Execute the flow with each data set.",
                   "Compare results against the rule definition."],
                  "Outcomes match the business rule for every data set, including edge combinations.",
                  "High")

        if re.search(r"\b(screen|page|button|display|ui|layout|menu|view|list)\b", text, re.I):
            b.add(module, req,
                  f"UI validation for: {_short(text, 90)}",
                  "UI Validation",
                  ["Open the screen referenced by the requirement.",
                   "Verify all elements, labels, and states described.",
                   "Resize the window / check on a second resolution."],
                  "UI matches the requirement; no truncation, overlap, or missing elements.",
                  "Low" if prio != "High" else "Medium")

        if re.search(r"\b(api|endpoint|service|integration|webhook|interface)\b", text, re.I):
            b.add(module, req,
                  f"API validation for: {_short(text, 90)}",
                  "API Validation",
                  ["Call the endpoint with a valid payload and auth token.",
                   "Call with an invalid payload and with missing auth.",
                   "Verify response codes, schema, and error bodies."],
                  "Valid call returns 2xx with the documented schema; invalid calls return 4xx with clear errors.",
                  "High")

    modules = sorted({r["module"] for r in requirements})
    for i in range(len(modules) - 1):
        a, c = modules[i], modules[i + 1]
        b.add(f"{a} + {c}", None,
              f"Integration — data flows correctly between {a} and {c}",
              "Integration",
              [f"Complete a transaction in {a}.",
               f"Verify the resulting state/data is reflected in {c}.",
               "Check audit/log entries for the handoff."],
              "Data is consistent across both modules with no manual intervention.",
              "High")

    if modules:
        b.add("End-to-End", None,
              "E2E — complete primary business flow across all modules: " + ", ".join(modules),
              "End-to-End",
              ["Start from a clean state as a new user.",
               "Execute the primary business journey touching each module in sequence.",
               "Verify final state, notifications, and reports."],
              "The full journey completes without errors and all module states are consistent.",
              "High",
              preconditions="Clean environment; all integrations up.")

    return {"test_cases": b.cases, "duplicates_skipped": b.duplicates_skipped,
            "duplicate_examples": b.duplicate_examples, "engine": "demo",
            "assumptions": _extract_assumptions(requirements)}


# ---------------------------------------------------------------------------
# AI-mode (Claude) test case generation
# ---------------------------------------------------------------------------

TC_SCHEMA = {
    "type": "object",
    "properties": {
        "test_cases": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "module": {"type": "string"},
                    "requirement_ref": {"type": "string"},
                    "requirement_description": {"type": "string"},
                    "user_story_ref": {"type": "string"},
                    "scenario": {"type": "string"},
                    "test_objective": {"type": "string"},
                    "preconditions": {"type": "string"},
                    "steps": {"type": "array", "items": {"type": "string"}},
                    "test_data": {"type": "string"},
                    "expected_result": {"type": "string"},
                    "priority": {"type": "string", "enum": ["High", "Medium", "Low"]},
                    "severity": {"type": "string", "enum": ["Critical", "High", "Medium", "Low"]},
                    "category": {"type": "string"},
                    "regression_impact": {"type": "string", "enum": ["High", "Medium", "Low"]},
                    "automation_candidate": {"type": "string", "enum": ["Yes", "No"]},
                    "traceability": {"type": "string"},
                },
                "required": ["module", "requirement_ref", "requirement_description",
                              "user_story_ref", "scenario", "test_objective", "preconditions",
                              "steps", "test_data", "expected_result", "priority", "severity",
                              "category", "regression_impact", "automation_candidate",
                              "traceability"],
                "additionalProperties": False,
            },
        },
        "assumptions": {
            "type": "array",
            "items": {"type": "string"},
            "description": "Assumptions made about unclear/underspecified requirements — "
                            "never mixed into the test cases themselves.",
        },
    },
    "required": ["test_cases", "assumptions"],
    "additionalProperties": False,
}

AI_SYSTEM = """You are a Senior Software QA Test Architect generating an enterprise-grade \
regression test suite from requirement documents (BRD/FRD/SRS/user stories).

Read and understand every requirement completely before generating anything. Identify: \
functional requirements, non-functional requirements, business rules, validation rules, \
dependencies, user roles, workflows, integrations, and edge cases. Never assume a feature \
that is not stated in the documents — if a requirement is vague or underspecified, record \
an assumption in the `assumptions` list instead of inventing behavior in a test case.

For every requirement, identify its happy path, alternate path, and exception path, and \
generate test cases covering: functional regression, positive, negative, boundary value, \
validation, integration, end-to-end, error handling/exception, business rule validation, \
data validation, UI validation, security, and API validation (where applicable). Every \
test case must trace to a requirement ID where possible, and steps must be concrete and \
executable by a manual tester.

Set `priority` by business risk (High for money/security/data-integrity paths) and \
`severity` by the impact of the underlying defect if the test fails (Critical for \
security/payment/compliance/data-integrity issues). Set `regression_impact` to how much \
the suite's regression coverage would degrade if this case were removed, and \
`automation_candidate` to whether a deterministic automated check is practical for it.

Maximize regression coverage with minimum redundancy — do not emit near-duplicate \
scenarios."""


def ai_available() -> bool:
    return bool(os.environ.get("ANTHROPIC_API_KEY") or os.environ.get("ANTHROPIC_AUTH_TOKEN"))


def ai_generate(requirements: list[dict], doc_texts: dict[str, str]) -> dict:
    import anthropic

    client = anthropic.Anthropic()
    combined = "\n\n".join(f"=== {name} ===\n{text[:40000]}" for name, text in doc_texts.items())
    user_msg = (
        "Extracted requirement candidates (use these IDs for traceability):\n"
        + json.dumps(requirements, indent=1)
        + "\n\nFull document text:\n" + combined[:120000]
        + "\n\nGenerate the complete regression test suite."
    )
    with client.messages.stream(
        model="claude-opus-5",
        max_tokens=64000,
        thinking={"type": "adaptive"},
        system=AI_SYSTEM,
        output_config={"format": {"type": "json_schema", "schema": TC_SCHEMA}},
        messages=[{"role": "user", "content": user_msg}],
    ) as stream:
        message = stream.get_final_message()

    text = next((b.text for b in message.content if b.type == "text"), "{}")
    parsed = json.loads(text)
    cases = parsed.get("test_cases", [])
    seen: set[str] = set()
    unique: list[dict] = []
    duplicates = 0
    duplicate_examples: list[str] = []
    for i, case in enumerate(cases, 1):
        norm = re.sub(r"[^a-z0-9]", "", case.get("scenario", "").lower())
        if norm in seen:
            duplicates += 1
            if len(duplicate_examples) < 25:
                duplicate_examples.append(case.get("scenario", ""))
            continue
        seen.add(norm)
        case["id"] = f"TC-{i:03d}"
        unique.append(case)
    return {"test_cases": unique, "duplicates_skipped": duplicates,
            "duplicate_examples": duplicate_examples, "engine": "ai",
            "assumptions": parsed.get("assumptions", [])}


# ---------------------------------------------------------------------------
# Coverage / report computation
# ---------------------------------------------------------------------------

def build_report(requirements: list[dict], test_cases: list[dict]) -> dict:
    req_ids = [r["id"] for r in requirements]
    tc_by_req: dict[str, list[str]] = {rid: [] for rid in req_ids}
    for tc in test_cases:
        ref = tc.get("requirement_ref", "")
        for rid in req_ids:
            if rid in ref:
                tc_by_req[rid].append(tc["id"])
                break

    covered = [rid for rid, tcs in tc_by_req.items() if tcs]
    thin = [rid for rid, tcs in tc_by_req.items() if 0 < len(tcs) < 2]
    uncovered = [rid for rid, tcs in tc_by_req.items() if not tcs]

    module_stats: dict[str, dict] = {}
    for req in requirements:
        stats = module_stats.setdefault(req["module"], {"requirements": 0, "test_cases": 0, "covered": 0})
        stats["requirements"] += 1
        if tc_by_req[req["id"]]:
            stats["covered"] += 1
    for tc in test_cases:
        module = tc.get("module", "General")
        stats = module_stats.setdefault(module, {"requirements": 0, "test_cases": 0, "covered": 0})
        stats["test_cases"] += 1

    priorities = {"High": 0, "Medium": 0, "Low": 0}
    categories: dict[str, int] = {}
    for tc in test_cases:
        priorities[tc.get("priority", "Medium")] = priorities.get(tc.get("priority", "Medium"), 0) + 1
        cat = tc.get("category", "Other")
        categories[cat] = categories.get(cat, 0) + 1

    return {
        "total_requirements": len(requirements),
        "total_test_cases": len(test_cases),
        "functional_count": sum(1 for r in requirements if r.get("type") != "Non-Functional"),
        "non_functional_count": sum(1 for r in requirements if r.get("type") == "Non-Functional"),
        "coverage_percent": round(100 * len(covered) / len(req_ids), 1) if req_ids else 0.0,
        "priorities": priorities,
        "categories": categories,
        "modules": module_stats,
        "traceability": [
            {
                "requirement_id": req["id"],
                "requirement": req["text"],
                "module": req["module"],
                "source": req["source"],
                "test_case_ids": tc_by_req[req["id"]],
                "covered": bool(tc_by_req[req["id"]]),
            }
            for req in requirements
        ],
        "alerts": (
            [f"{rid}: no test cases generated — review this requirement manually." for rid in uncovered]
            + [f"{rid}: thin coverage (only 1 test case)." for rid in thin]
        ),
    }


# ---------------------------------------------------------------------------
# Risk scoring
# ---------------------------------------------------------------------------

def compute_risk(report: dict, business_rules: list[dict]) -> dict:
    total_req = report["total_requirements"] or 1
    uncovered = sum(1 for t in report["traceability"] if not t["covered"])
    uncovered_pct = 100 * uncovered / total_req
    high_pct = 100 * report["priorities"].get("High", 0) / max(report["total_test_cases"], 1)
    rule_density = min(100, 100 * len(business_rules) / total_req)
    score = round(0.45 * uncovered_pct + 0.35 * high_pct + 0.20 * rule_density, 1)
    if score >= 65:
        level = "Critical"
    elif score >= 45:
        level = "High"
    elif score >= 22:
        level = "Medium"
    else:
        level = "Low"
    return {"score": score, "level": level}


# ---------------------------------------------------------------------------
# AI insights
# ---------------------------------------------------------------------------

def build_insights(test_cases: list[dict], report: dict, business_rules: list[dict],
                    duplicate_examples: list[str]) -> dict:
    modules = report["modules"]

    high_risk_modules = [m for m, s in modules.items() if s["requirements"] and s["covered"] < s["requirements"]]
    tc_high_by_module: dict[str, int] = {}
    for tc in test_cases:
        if tc.get("priority") == "High":
            tc_high_by_module[tc["module"]] = tc_high_by_module.get(tc["module"], 0) + 1
    for m, s in modules.items():
        if m in high_risk_modules:
            continue
        if s["test_cases"] and tc_high_by_module.get(m, 0) / s["test_cases"] >= 0.6:
            high_risk_modules.append(m)

    uncovered = [t for t in report["traceability"] if not t["covered"]]
    thin = [t for t in report["traceability"] if 0 < len(t["test_case_ids"]) < 2]
    recommended_areas = [m for m, _ in sorted(modules.items(), key=lambda kv: kv[1]["test_cases"], reverse=True)[:3]]

    return {
        "high_risk_modules": high_risk_modules,
        "missing_coverage": [{"id": t["requirement_id"], "text": t["requirement"]} for t in uncovered],
        "duplicate_test_cases": duplicate_examples,
        "untested_requirements": [{"id": t["requirement_id"], "text": t["requirement"]} for t in uncovered],
        "recommended_regression_areas": recommended_areas,
        "suggested_additional_test_cases": [
            {"id": t["requirement_id"], "text": t["requirement"],
             "suggestion": "Only thin coverage exists today — add a negative or boundary-value case."}
            for t in thin
        ],
        "business_rule_conflicts": find_business_rule_conflicts(business_rules),
        "requirement_gaps": [{"id": t["requirement_id"], "text": t["requirement"]} for t in uncovered],
    }


# ---------------------------------------------------------------------------
# Impact analysis (diff against the previous analysis of this project)
# ---------------------------------------------------------------------------

def compute_impact(previous_requirements: list[dict] | None, requirements: list[dict],
                    test_cases: list[dict]) -> dict | None:
    if not previous_requirements:
        return None

    def norm(t: str) -> str:
        return re.sub(r"[^a-z0-9]", "", t.lower())

    prev = {norm(r["text"]): r for r in previous_requirements}
    curr = {norm(r["text"]): r for r in requirements}
    added = [r for k, r in curr.items() if k not in prev]
    removed = [r for k, r in prev.items() if k not in curr]

    impacted_modules = sorted({r["module"] for r in added} | {r["module"] for r in removed})
    impacted_test_cases = [tc["id"] for tc in test_cases if tc.get("module") in impacted_modules]
    high_risk_added = [r for r in added if HIGH_PRIORITY.search(r["text"])]
    risk_level = "High" if high_risk_added else ("Medium" if (added or removed) else "Low")
    order = sorted(
        impacted_modules,
        key=lambda m: (0 if any(r["module"] == m for r in high_risk_added) else 1, m),
    )

    return {
        "added": [{"id": r["id"], "text": r["text"], "module": r["module"]} for r in added],
        "removed": [{"id": r["id"], "text": r["text"], "module": r["module"]} for r in removed],
        "impacted_modules": impacted_modules,
        "impacted_test_cases": impacted_test_cases,
        "risk_level": risk_level,
        "recommended_execution_order": order,
    }


# ---------------------------------------------------------------------------
# AI chat assistant
# ---------------------------------------------------------------------------

def rule_based_answer(question: str, project: dict) -> str:
    q = question.lower()
    report = project.get("report") or {}
    insights = project.get("insights") or {}
    risk = project.get("risk") or {}
    business_rules = project.get("business_rules") or []
    impact = project.get("impact")
    test_cases = project.get("test_cases") or []

    m = re.search(r"\bTC-\d+\b", question, re.I)
    if m:
        tc = next((t for t in test_cases if t["id"].upper() == m.group(0).upper()), None)
        if tc:
            return (f"{tc['id']} ({tc['module']}) — {tc['scenario']}\n"
                    f"Traces to {tc['requirement_ref']}. Category: {tc['category']}, Priority: {tc['priority']}.\n"
                    f"Expected result: {tc['expected_result']}")
        return f"I couldn't find test case {m.group(0)} in this project."

    m = re.search(r"\bBR-\d+\b", question, re.I)
    if m:
        rule = next((r for r in business_rules if r["id"].upper() == m.group(0).upper()), None)
        if rule:
            return f"{rule['id']} ({rule['module']}): {rule['text']}\nTraces to {rule['requirement_ref']} in {rule['source']}."
        return f"I couldn't find business rule {m.group(0)} — {len(business_rules)} business rule(s) detected in this project."

    if "impact" in q or "which module" in q or "impacted" in q:
        if impact:
            return ("Since the last analysis, these modules are impacted: "
                     + ", ".join(impact["impacted_modules"] or ["none"])
                     + f". Risk level: {impact['risk_level']}. Recommended execution order: "
                     + ", ".join(impact["recommended_execution_order"] or ["n/a"]) + ".")
        modules = sorted(report.get("modules", {}).keys())
        return ("No prior version to compare against yet, so nothing is 'newly' impacted. "
                 "All modules currently in scope: " + ", ".join(modules or ["none"]) + ".")

    if "high-priority" in q or "high priority" in q:
        cases = [t for t in test_cases if t["priority"] == "High"]
        preview = "; ".join(t["id"] for t in cases[:10])
        return f"{len(cases)} High-priority regression test case(s). First {min(10, len(cases))}: {preview or 'none'}."

    if "compar" in q and "version" in q:
        if impact:
            return (f"Compared to the previous analysis: {len(impact['added'])} requirement(s) added, "
                     f"{len(impact['removed'])} removed, impacting {len(impact['impacted_modules'])} module(s) "
                     f"and {len(impact['impacted_test_cases'])} test case(s). Risk: {impact['risk_level']}.")
        return "This project has only been analyzed once, so there's no prior version to compare against."

    if "duplicate" in q:
        dups = insights.get("duplicate_test_cases", [])
        extra = (" Examples: " + "; ".join(dups[:5])) if dups else ""
        return f"{len(dups)} duplicate scenario(s) were detected and skipped during generation.{extra}"

    if "recommend" in q or "additional" in q or "suggest" in q:
        sugg = insights.get("suggested_additional_test_cases", [])
        if not sugg:
            return "Coverage looks solid — no requirement currently has thin coverage."
        preview = "; ".join(s["id"] for s in sugg[:8])
        return f"{len(sugg)} requirement(s) have thin coverage and would benefit from an additional case: {preview}."

    if "risk" in q:
        return (f"Overall risk score: {risk.get('score', 'n/a')}/100 ({risk.get('level', 'n/a')}). "
                 "Based on requirement coverage gaps, high-priority test density, and business rule volume.")

    return (f"This project has {report.get('total_requirements', 0)} requirement(s), "
             f"{report.get('total_test_cases', 0)} test case(s), {report.get('coverage_percent', 0)}% coverage, "
             f"and an overall risk level of {risk.get('level', 'n/a')}. Ask about a specific TC-/BR- id, "
             "high-priority tests, impacted modules, duplicates, or recommended additional coverage.")


def ai_answer(question: str, project: dict) -> str:
    import anthropic

    client = anthropic.Anthropic()
    context = {
        "report": project.get("report"),
        "insights": project.get("insights"),
        "risk": project.get("risk"),
        "business_rules": project.get("business_rules"),
        "impact": project.get("impact"),
        "test_cases": (project.get("test_cases") or [])[:60],
    }
    msg = client.messages.create(
        model="claude-opus-5",
        max_tokens=600,
        system=("You are a QA assistant answering questions about a regression test suite. "
                 "Use only the JSON context provided — do not invent requirement or test case IDs. "
                 "Be concise (under 120 words)."),
        messages=[{"role": "user", "content": f"Context:\n{json.dumps(context)[:60000]}\n\nQuestion: {question}"}],
    )
    return next((b.text for b in msg.content if b.type == "text"), "").strip()


def answer_question(question: str, project: dict) -> str:
    if ai_available():
        try:
            answer = ai_answer(question, project)
            if answer:
                return answer
        except Exception:  # noqa: BLE001 — fall back to rule-based answers
            pass
    return rule_based_answer(question, project)
