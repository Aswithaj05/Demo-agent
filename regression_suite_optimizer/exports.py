"""Export generated test cases to CSV, Excel (XLSX), Word (DOCX), and PDF.

XLSX and PDF are produced with lightweight built-in writers (no openpyxl /
reportlab dependency); DOCX uses python-docx; CSV uses the stdlib.
"""

from __future__ import annotations

import csv
import io
import zipfile
from xml.sax.saxutils import escape

COLUMNS = [
    ("id", "Test Case ID"),
    ("module", "Module"),
    ("requirement_ref", "Requirement ID"),
    ("requirement_description", "Requirement Description"),
    ("user_story_ref", "User Story Ref"),
    ("scenario", "Test Scenario"),
    ("test_objective", "Test Objective"),
    ("preconditions", "Preconditions"),
    ("steps", "Test Steps"),
    ("test_data", "Test Data"),
    ("expected_result", "Expected Result"),
    ("priority", "Priority"),
    ("severity", "Severity"),
    ("category", "Test Type"),
    ("regression_impact", "Regression Impact"),
    ("automation_candidate", "Automation Candidate"),
    ("traceability", "Traceability"),
]


def _cell(tc: dict, key: str) -> str:
    value = tc.get(key, "")
    if key == "steps" and isinstance(value, list):
        return "\n".join(f"{i}. {s}" for i, s in enumerate(value, 1))
    return str(value)


# ---------------------------------------------------------------------------
# CSV
# ---------------------------------------------------------------------------

def to_csv(test_cases: list[dict]) -> bytes:
    buf = io.StringIO()
    writer = csv.writer(buf, lineterminator="\r\n")
    writer.writerow([header for _, header in COLUMNS])
    for tc in test_cases:
        writer.writerow([_cell(tc, key) for key, _ in COLUMNS])
    return buf.getvalue().encode("utf-8-sig")  # BOM so Excel opens it cleanly


# ---------------------------------------------------------------------------
# XLSX (minimal OOXML writer, inline strings)
# ---------------------------------------------------------------------------

_XLSX_CONTENT_TYPES = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
<Default Extension="xml" ContentType="application/xml"/>
<Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>
<Override PartName="/xl/worksheets/sheet1.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>
<Override PartName="/xl/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.styles+xml"/>
</Types>"""

_XLSX_RELS = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>
</Relationships>"""

_XLSX_WORKBOOK = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
<sheets><sheet name="Regression Suite" sheetId="1" r:id="rId1"/></sheets>
</workbook>"""

_XLSX_WORKBOOK_RELS = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet1.xml"/>
<Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>
</Relationships>"""

_XLSX_STYLES = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<styleSheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
<fonts count="2"><font><sz val="11"/><name val="Calibri"/></font><font><b/><sz val="11"/><name val="Calibri"/></font></fonts>
<fills count="2"><fill><patternFill patternType="none"/></fill><fill><patternFill patternType="gray125"/></fill></fills>
<borders count="1"><border/></borders>
<cellStyleXfs count="1"><xf/></cellStyleXfs>
<cellXfs count="2"><xf/><xf fontId="1" applyFont="1"/></cellXfs>
</styleSheet>"""


def _xlsx_row(values: list[str], row_num: int, style: int = 0) -> str:
    cells = []
    for value in values:
        text = escape(value)[:32000]
        style_attr = f' s="{style}"' if style else ""
        cells.append(f'<c t="inlineStr"{style_attr}><is><t xml:space="preserve">{text}</t></is></c>')
    return f'<row r="{row_num}">{"".join(cells)}</row>'


def to_xlsx(test_cases: list[dict]) -> bytes:
    rows = [_xlsx_row([header for _, header in COLUMNS], 1, style=1)]
    for i, tc in enumerate(test_cases, 2):
        rows.append(_xlsx_row([_cell(tc, key) for key, _ in COLUMNS], i))
    sheet = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">'
        "<sheetData>" + "".join(rows) + "</sheetData></worksheet>"
    )
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", _XLSX_CONTENT_TYPES)
        zf.writestr("_rels/.rels", _XLSX_RELS)
        zf.writestr("xl/workbook.xml", _XLSX_WORKBOOK)
        zf.writestr("xl/_rels/workbook.xml.rels", _XLSX_WORKBOOK_RELS)
        zf.writestr("xl/styles.xml", _XLSX_STYLES)
        zf.writestr("xl/worksheets/sheet1.xml", sheet)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# DOCX (python-docx)
# ---------------------------------------------------------------------------

def to_docx(test_cases: list[dict], project_name: str) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading("Regression Test Suite", level=0)
    doc.add_paragraph(f"Project: {project_name}")
    doc.add_paragraph(f"Total test cases: {len(test_cases)}")

    for tc in test_cases:
        doc.add_heading(f"{tc.get('id', '')} — {tc.get('scenario', '')}", level=2)
        table = doc.add_table(rows=0, cols=2)
        table.style = "Light Grid Accent 1"
        for key, header in COLUMNS:
            if key in ("id", "scenario"):
                continue
            row = table.add_row().cells
            row[0].text = header
            row[1].text = _cell(tc, key)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF (minimal text-based writer, Helvetica)
# ---------------------------------------------------------------------------

_PAGE_W, _PAGE_H = 612, 792  # US Letter, points
_MARGIN, _LEAD = 46, 12.5
_LINES_PER_PAGE = int((_PAGE_H - 2 * _MARGIN) / _LEAD)
_WRAP = 100


def _pdf_escape(text: str) -> str:
    text = text.encode("latin-1", errors="replace").decode("latin-1")
    return text.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")


def _wrap_text(text: str, width: int) -> list[str]:
    out: list[str] = []
    for raw in text.split("\n"):
        line = ""
        for word in raw.split(" "):
            candidate = f"{line} {word}".strip()
            if len(candidate) <= width:
                line = candidate
            else:
                if line:
                    out.append(line)
                while len(word) > width:
                    out.append(word[:width])
                    word = word[width:]
                line = word
        out.append(line)
    return out


def to_pdf(test_cases: list[dict], project_name: str) -> bytes:
    lines: list[tuple[str, bool]] = [
        (f"Regression Test Suite — {project_name}", True),
        (f"Total test cases: {len(test_cases)}", False),
        ("", False),
    ]
    for tc in test_cases:
        lines.append((f"{tc.get('id', '')}  [{tc.get('priority', '')}]  {tc.get('category', '')}", True))
        for label, key in [("Module", "module"), ("Requirement", "requirement_ref"),
                            ("Requirement Description", "requirement_description"),
                            ("Scenario", "scenario"), ("Test Objective", "test_objective"),
                            ("Preconditions", "preconditions"),
                            ("Steps", "steps"), ("Test Data", "test_data"),
                            ("Expected", "expected_result"), ("Severity", "severity"),
                            ("Regression Impact", "regression_impact"),
                            ("Automation Candidate", "automation_candidate"),
                            ("Traceability", "traceability")]:
            for wrapped in _wrap_text(f"{label}: {_cell(tc, key)}", _WRAP):
                lines.append((wrapped, False))
        lines.append(("", False))
    return _lines_to_pdf(lines)


def _lines_to_pdf(lines: list[tuple[str, bool]]) -> bytes:
    """Render a list of (text, bold) lines to a minimal multi-page PDF."""
    pages = [lines[i:i + _LINES_PER_PAGE] for i in range(0, len(lines), _LINES_PER_PAGE)] or [[]]

    objects: list[bytes] = []

    def add(obj: bytes) -> int:
        objects.append(obj)
        return len(objects)  # 1-based object number

    font_regular = add(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    font_bold = add(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica-Bold >>")

    page_object_numbers: list[int] = []
    content_object_numbers: list[int] = []
    for page in pages:
        parts = ["BT", f"1 0 0 1 {_MARGIN} {_PAGE_H - _MARGIN} Tm", f"{_LEAD} TL"]
        current_bold = None
        for text, bold in page:
            if bold != current_bold:
                parts.append(f"/{'F2' if bold else 'F1'} {'10' if bold else '9'} Tf")
                current_bold = bold
            parts.append(f"({_pdf_escape(text)}) Tj T*")
        parts.append("ET")
        stream = "\n".join(parts).encode("latin-1")
        content_object_numbers.append(
            add(b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream")
        )
        page_object_numbers.append(0)  # placeholder, filled after pages object exists

    pages_obj_num = len(objects) + len(pages) + 1
    for i, content_num in enumerate(content_object_numbers):
        page_object_numbers[i] = add(
            (f"<< /Type /Page /Parent {pages_obj_num} 0 R /MediaBox [0 0 {_PAGE_W} {_PAGE_H}] "
             f"/Resources << /Font << /F1 {font_regular} 0 R /F2 {font_bold} 0 R >> >> "
             f"/Contents {content_num} 0 R >>").encode()
        )
    kids = " ".join(f"{n} 0 R" for n in page_object_numbers)
    pages_num = add(f"<< /Type /Pages /Kids [{kids}] /Count {len(pages)} >>".encode())
    assert pages_num == pages_obj_num
    catalog_num = add(f"<< /Type /Catalog /Pages {pages_num} 0 R >>".encode())

    out = io.BytesIO()
    out.write(b"%PDF-1.4\n")
    offsets = [0]
    for i, obj in enumerate(objects, 1):
        offsets.append(out.tell())
        out.write(f"{i} 0 obj\n".encode() + obj + b"\nendobj\n")
    xref_pos = out.tell()
    out.write(f"xref\n0 {len(objects) + 1}\n".encode())
    out.write(b"0000000000 65535 f \n")
    for off in offsets[1:]:
        out.write(f"{off:010d} 00000 n \n".encode())
    out.write(
        f"trailer\n<< /Size {len(objects) + 1} /Root {catalog_num} 0 R >>\n"
        f"startxref\n{xref_pos}\n%%EOF".encode()
    )
    return out.getvalue()


# ---------------------------------------------------------------------------
# Additional report exports: traceability matrix, coverage / impact / summary
# ---------------------------------------------------------------------------

def to_traceability_csv(report: dict) -> bytes:
    buf = io.StringIO()
    writer = csv.writer(buf, lineterminator="\r\n")
    writer.writerow(["Requirement ID", "Requirement", "Module", "Source", "Test Case IDs", "Coverage %", "Status"])
    for t in report.get("traceability", []):
        writer.writerow([
            t["requirement_id"], t["requirement"], t["module"], t["source"],
            ", ".join(t["test_case_ids"]),
            "100" if t["covered"] else "0",
            "Covered" if t["covered"] else "Gap",
        ])
    return buf.getvalue().encode("utf-8-sig")


def to_coverage_pdf(report: dict, project_name: str) -> bytes:
    lines: list[tuple[str, bool]] = [
        (f"Coverage Report — {project_name}", True),
        (f"Requirements: {report['total_requirements']}  |  Test cases: {report['total_test_cases']}  |  "
         f"Coverage: {report['coverage_percent']}%", False),
        ("", False),
        ("Module-wise coverage", True),
    ]
    for module, s in report.get("modules", {}).items():
        pct = round(100 * s["covered"] / s["requirements"], 1) if s["requirements"] else 100.0
        lines.append((f"{module}: {s['requirements']} req(s), {s['test_cases']} test(s), {pct}% covered", False))
    lines.append(("", False))
    lines.append(("Alerts", True))
    for a in report.get("alerts", []) or ["None — all requirements have coverage."]:
        lines.append((a, False))
    return _lines_to_pdf(lines)


def to_impact_pdf(impact: dict | None, project_name: str) -> bytes:
    lines: list[tuple[str, bool]] = [(f"Impact Analysis — {project_name}", True), ("", False)]
    if not impact:
        lines.append(("This project has only been analyzed once — no prior version to compare against.", False))
        return _lines_to_pdf(lines)
    lines.append((f"Risk level: {impact['risk_level']}", False))
    lines.append((f"Impacted modules: {', '.join(impact['impacted_modules']) or 'none'}", False))
    lines.append((f"Impacted test cases: {', '.join(impact['impacted_test_cases']) or 'none'}", False))
    lines.append((f"Recommended execution order: {', '.join(impact['recommended_execution_order']) or 'n/a'}", False))
    lines.append(("", False))
    lines.append((f"Added requirements ({len(impact['added'])})", True))
    for r in impact["added"]:
        lines.append((f"{r['id']} [{r['module']}]: {r['text']}", False))
    lines.append(("", False))
    lines.append((f"Removed requirements ({len(impact['removed'])})", True))
    for r in impact["removed"]:
        lines.append((f"{r['id']} [{r['module']}]: {r['text']}", False))
    return _lines_to_pdf(lines)


def to_summary_pdf(report: dict, insights: dict, risk: dict, project_name: str,
                    assumptions: list[str] | None = None) -> bytes:
    lines: list[tuple[str, bool]] = [
        (f"Regression Summary — {project_name}", True),
        (f"Requirements: {report['total_requirements']} (Functional {report.get('functional_count', 0)}, "
         f"Non-Functional {report.get('non_functional_count', 0)})", False),
        (f"Test cases: {report['total_test_cases']}  |  Coverage: {report['coverage_percent']}%  |  "
         f"Risk: {risk.get('score', 'n/a')}/100 ({risk.get('level', 'n/a')})", False),
        (f"Priority split — High: {report['priorities'].get('High', 0)}  "
         f"Medium: {report['priorities'].get('Medium', 0)}  Low: {report['priorities'].get('Low', 0)}", False),
        ("", False),
        ("High-risk modules", True),
    ]
    lines += [(m, False) for m in (insights.get("high_risk_modules") or ["None identified."])]
    lines.append(("", False))
    lines.append(("Recommended regression areas", True))
    lines += [(m, False) for m in (insights.get("recommended_regression_areas") or ["None."])]
    lines.append(("", False))
    lines.append(("Requirement gaps", True))
    lines += [(f"{g['id']}: {g['text']}", False) for g in (insights.get("requirement_gaps") or [])] or [("None.", False)]
    lines.append(("", False))
    lines.append(("Business rule conflicts", True))
    lines += [(c["note"], False) for c in (insights.get("business_rule_conflicts") or [])] or [("None detected.", False)]
    lines.append(("", False))
    lines.append(("Assumptions (requirements were unclear or underspecified)", True))
    lines += [(a, False) for a in (assumptions or [])] or [("None — all requirements were unambiguous.", False)]
    return _lines_to_pdf(lines)


REPORT_EXPORTERS = {
    "traceability": (to_traceability_csv, "text/csv", "csv"),
    "coverage": (to_coverage_pdf, "application/pdf", "pdf"),
    "impact": (to_impact_pdf, "application/pdf", "pdf"),
    "summary": (to_summary_pdf, "application/pdf", "pdf"),
}


EXPORTERS = {
    "csv": (to_csv, "text/csv", "csv"),
    "xlsx": (to_xlsx, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "xlsx"),
    "docx": (to_docx, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"),
    "pdf": (to_pdf, "application/pdf", "pdf"),
}
